import operator
import subprocess
import os
import shutil
import glob
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT

document = Document()
#section = document.sections[-1]
sections = document.sections
for section in document.sections:
  section.left_margin = Inches(1)
  section.right_margin = Inches(1)
  section.orientation = WD_ORIENT.LANDSCAPE
  new_width, new_height = section.page_height, section.page_width
  section.page_width = new_width
  section.page_height = new_height

files = sorted(glob.glob ('*.png'))
fignum = 1
for fname in files:
  document.add_picture(fname, width=Inches(9)) # width=Inches(6.5))
  document.add_paragraph('Figure ' + str(fignum) + ': ' + fname, style='Caption')
#  document.add_page_break()
  fignum += 1

document.save('plots.docx')

